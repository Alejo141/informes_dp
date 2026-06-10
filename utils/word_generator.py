"""
utils/word_generator.py
Genera el documento Word final a partir de:
  - plantilla base (docx del mes anterior)
  - métricas procesadas
  - gráficos generados
  - configuración del cliente
  - capítulos seleccionados
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re
from pathlib import Path

MESES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

AZUL = RGBColor(0x1B, 0x4F, 0x8A)
GRIS = RGBColor(0xF2, 0xF2, 0xF2)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)


def fmt_num(n):
    """Formatea número con puntos de miles."""
    try:
        return f"{int(n):,}".replace(",", ".")
    except:
        return str(n)


def fmt_cop(n):
    """Formatea valor en pesos colombianos."""
    try:
        return f"$ {int(n):,}".replace(",", ".")
    except:
        return str(n)


class WordGenerator:
    def __init__(self, config: dict, mes: int, ano: int,
                 template_path: str, metricas: dict,
                 graficos: dict, capitulos_seleccionados: list):
        self.config = config
        self.mes = mes
        self.ano = ano
        self.mes_nombre = MESES.get(mes, str(mes))
        self.template_path = template_path
        self.metricas = metricas
        self.graficos = graficos
        self.capitulos_sel = set(capitulos_seleccionados)  # set de IDs

        # Cargar documento base
        self.doc = Document(template_path) if template_path else Document()
        self._setup_styles()

    # ------------------------------------------------------------------ #
    # Estilos
    # ------------------------------------------------------------------ #
    def _setup_styles(self):
        """Asegura que los estilos de heading tengan el formato correcto."""
        from docx.oxml.ns import qn
        try:
            s = self.doc.styles['Heading 1']
            s.font.size = Pt(13)
            s.font.bold = True
            s.font.color.rgb = AZUL
        except:
            pass
        try:
            s = self.doc.styles['Heading 2']
            s.font.size = Pt(11)
            s.font.bold = True
            s.font.color.rgb = AZUL
        except:
            pass

    # ------------------------------------------------------------------ #
    # Helpers de inserción
    # ------------------------------------------------------------------ #
    def _add_heading(self, text: str, level: int = 1):
        p = self.doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = AZUL
        run.font.bold = True
        return p

    def _add_paragraph(self, text: str, bold=False, italic=False, size=10):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def _add_image(self, path: str, width_inches: float = 6.0):
        if not path or not Path(path).exists():
            return
        try:
            self.doc.add_picture(path, width=Inches(width_inches))
            last_par = self.doc.paragraphs[-1]
            last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            self._add_paragraph(f"[Gráfico no disponible: {e}]", italic=True)

    def _set_cell_bg(self, cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _build_table(self, headers: list, rows: list, col_widths: list = None):
        """Crea una tabla formateada."""
        ncols = len(headers)
        table = self.doc.add_table(rows=1, cols=ncols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Cabecera
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = str(h)
            self._set_cell_bg(cell, "1B4F8A")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = BLANCO
                run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Filas de datos
        for ridx, row_data in enumerate(rows):
            row = table.add_row()
            bg = "F2F2F2" if ridx % 2 == 0 else "FFFFFF"
            for cidx, val in enumerate(row_data):
                cell = row.cells[cidx]
                cell.text = str(val) if val is not None else ""
                self._set_cell_bg(cell, bg)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Anchos de columna
        if col_widths:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tbl = table._tbl
            for i, row in enumerate(tbl.findall(qn('w:tr'))):
                for j, tc in enumerate(row.findall(qn('w:tc'))):
                    if j < len(col_widths):
                        tcPr = tc.get_or_add_tcPr()
                        w_el = OxmlElement('w:tcW')
                        w_el.set(qn('w:w'), str(col_widths[j]))
                        w_el.set(qn('w:type'), 'dxa')
                        tcPr.append(w_el)
        return table

    # ------------------------------------------------------------------ #
    # Reemplazo de texto en plantilla
    # ------------------------------------------------------------------ #
    def _replace_in_template(self):
        """Reemplaza mes/año/cliente en el documento base."""
        mes_ant = MESES.get(self.mes - 1, MESES[12]) if self.mes > 1 else MESES[12]
        replacements = {
            # Patrones comunes de mes anterior → mes actual
            mes_ant.upper(): self.mes_nombre.upper(),
            mes_ant.capitalize(): self.mes_nombre.capitalize(),
        }

        for para in self.doc.paragraphs:
            for old, new in replacements.items():
                if old in para.text:
                    for run in para.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            if old in para.text:
                                for run in para.runs:
                                    if old in run.text:
                                        run.text = run.text.replace(old, new)

    # ------------------------------------------------------------------ #
    # Generador principal
    # ------------------------------------------------------------------ #
    def generar(self, output_path: str):
        """Genera el documento Word completo."""
        self._replace_in_template()

        # Limpiar contenido existente del body (excepto header/footer)
        # y agregar las secciones seleccionadas
        # Estrategia: agregar al final del documento existente
        self.doc.add_page_break()

        all_caps = self.config.get("capitulos", [])
        selected = self._get_selected_with_parents(all_caps)

        for cap in all_caps:
            if cap["id"] not in selected:
                continue
            num = cap["num"]
            titulo = cap["titulo"].upper()
            level_depth = len(num.split("."))

            if level_depth == 1:
                self._add_heading(f"{num}. {titulo}", level=1)
            elif level_depth == 2:
                self._add_heading(f"{num}. {titulo}", level=2)
            else:
                self._add_heading(f"{num}. {titulo}", level=3)

            # Dispatcher: genera contenido por capítulo
            self._render_capitulo(cap["id"], cap["num"])

        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path_obj))
        return str(output_path_obj)

    def _get_selected_with_parents(self, all_caps: list) -> set:
        """Si se selecciona un capítulo, incluye sus padres automáticamente."""
        sel = set(self.capitulos_sel)
        id_to_cap = {c["id"]: c for c in all_caps}
        to_add = set()
        for cid in sel:
            cap = id_to_cap.get(cid)
            if cap:
                parent_id = cap.get("parent")
                while parent_id:
                    to_add.add(parent_id)
                    parent_cap = id_to_cap.get(parent_id)
                    parent_id = parent_cap.get("parent") if parent_cap else None
        return sel | to_add

    def _render_capitulo(self, cap_id: int, num: str):
        """Dispatch a la función de renderizado según el ID del capítulo."""
        dispatch = {
            1:  self._cap_info_basica,
            2:  self._cap_descripcion,
            3:  self._cap_caracteristicas,
            4:  self._cap_obj_general,
            5:  self._cap_obj_especifico,
            6:  self._cap_localizacion,
            7:  self._cap_actividades_ejecutadas,
            8:  self._cap_gestion_operaciones,
            9:  self._cap_reposiciones,
            10: self._cap_gestion_comercial,
            11: self._cap_gestion_sedes,
            12: self._cap_gestion_sac,
            13: self._cap_gestion_social,
            14: self._cap_subsidios,
            15: self._cap_recaudo,
            16: self._cap_cartera,
            17: self._cap_financiera,
            18: self._cap_balance,
            19: self._cap_tarifas,
            20: self._cap_hseq,
            21: self._cap_sst,
            22: self._cap_control_personal,
            23: self._cap_ambiental,
            24: self._cap_actas,
            25: self._cap_juridico,
            26: self._cap_perfeccionamiento,
            27: self._cap_garantias,
            28: self._cap_polizas_rce,
            29: self._cap_poliza_todo_riesgo,
            30: self._cap_fiducia,
            31: self._cap_giros_fiducia,
            32: self._cap_distribucion_saldo,
        }
        fn = dispatch.get(cap_id)
        if fn:
            fn()

    # ------------------------------------------------------------------ #
    # CAPÍTULOS
    # ------------------------------------------------------------------ #
    def _cap_info_basica(self):
        headers = ["Campo", "Valor"]
        rows = [
            ["Contrato de concesión AMGC", self.config.get("contrato", "")],
            ["Supervisor del contrato", self.config.get("nombre_completo", "")],
            ["Interventor del contrato", self.config.get("interventor", "")],
            ["Ejecutor", f'{self.config.get("ejecutor", "")} NIT {self.config.get("nit_ejecutor", "")}'],
        ]
        self._add_paragraph(f"Tabla 1. Información Básica – {self.mes_nombre} {self.ano}", bold=True)
        self._build_table(headers, rows, [3000, 6000])
        self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

    def _cap_descripcion(self):
        self._add_paragraph(
            f"La Concesión de la operación técnica, tecnológica, logística, jurídica, contable, regulatoria, "
            f"predial, ambiental y demás actividades AOM de los sistemas SISFV en el marco del contrato "
            f"{self.config.get('contrato','')} ejecutado por {self.config.get('ejecutor','')} "
            f"para el periodo {self.mes_nombre} {self.ano}."
        )

    def _cap_caracteristicas(self):
        self._add_paragraph(
            f"{self.config.get('ejecutor','')} se compromete a desarrollar todas las actividades de "
            "Administración Mantenimiento y Gestión Comercial (AMGC) durante un plazo de diez (10) años, "
            "dando cumplimiento al marco regulatorio vigente, así como a lo establecido en el contrato de concesión."
        )

    def _cap_obj_general(self):
        self._add_paragraph(
            f"Presentar las actividades ejecutadas durante el mes de {self.mes_nombre} de {self.ano} "
            "en el marco de cumplimiento del Contrato de Concesión AOM."
        )

    def _cap_obj_especifico(self):
        bullets = [
            "Soportar la ejecución mensual de cada una de las actividades contempladas en el contrato de concesión.",
            "Cumplir con la obligación contractual de entrega de informes mensuales.",
            "Suministrar información a la supervisión e interventoría sobre las actividades AMGC de los sistemas SISFV."
        ]
        for b in bullets:
            p = self.doc.add_paragraph()
            p.add_run(f"• {b}").font.size = Pt(10)

    def _cap_localizacion(self):
        self._add_paragraph(
            f"El presente informe se enfoca en la intervención en zonas de difícil acceso, alejadas de las "
            "cabeceras municipales, en diferentes municipios pertenecientes a once (11) departamentos del país."
        )

    def _cap_actividades_ejecutadas(self):
        self._add_paragraph(
            f"Dando cumplimiento a las actividades contractuales, a continuación se presenta el estado de gestión "
            f"de los proyectos en ejecución durante {self.mes_nombre} {self.ano}."
        )
        ops = self.metricas.get("operaciones", {})
        por_mun = ops.get("por_municipio", {})
        if por_mun:
            headers = ["N°", "Municipio/Proyecto", "Visitas Funcional", "No Funcional", "Total"]
            rows = [[i+1, mun, d["funcional"], d["no_funcional"], d["total"]]
                    for i, (mun, d) in enumerate(por_mun.items())]
            totals = ["", "TOTAL",
                      sum(d["funcional"] for d in por_mun.values()),
                      sum(d["no_funcional"] for d in por_mun.values()),
                      sum(d["total"] for d in por_mun.values())]
            rows.append(totals)
            self._add_paragraph(f"Tabla. Estado de proyectos {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

    def _cap_gestion_operaciones(self):
        ops = self.metricas.get("operaciones", {})
        total = ops.get("total", 0)
        prev = ops.get("preventivos", 0)
        corr = ops.get("correctivos", 0)
        func = ops.get("funcionales", 0)
        no_func = ops.get("no_funcionales", 0)

        self._add_paragraph(
            f"Durante el mes de {self.mes_nombre} de {self.ano}, se ejecutaron un total de "
            f"{fmt_num(total)} visitas de mantenimiento, distribuidas así: {fmt_num(prev)} mantenimientos "
            f"preventivos y {fmt_num(corr)} mantenimientos correctivos."
        )

        if total > 0:
            headers = ["Tipo de Mantenimiento", "Cantidad"]
            rows = [["Mantenimientos correctivos", fmt_num(corr)],
                    ["Mantenimientos preventivos", fmt_num(prev)],
                    ["Total", fmt_num(total)]]
            self._add_paragraph(f"Tabla. Total mantenimientos ejecutados {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows, [5000, 3000])
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

            headers2 = ["Funcionalidad", "Cantidad"]
            rows2 = [["SISFV funcional", fmt_num(func)],
                     ["No funcionales", fmt_num(no_func)],
                     ["Total", fmt_num(total)]]
            self._add_paragraph(f"Tabla. Funcionalidad de visitas {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers2, rows2, [5000, 3000])
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        # Gráfico por municipio
        g = self.graficos.get("operaciones_municipio")
        if g:
            self._add_paragraph(f"Gráfico. Estado del parque por municipio – {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g, 6.5)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        # Tabla por municipio
        por_mun = ops.get("por_municipio", {})
        if por_mun:
            headers3 = ["Municipio", "Funcional", "No Funcional", "Total"]
            rows3 = [[mun, d["funcional"], d["no_funcional"], d["total"]]
                     for mun, d in por_mun.items()]
            totals3 = ["TOTAL",
                       sum(d["funcional"] for d in por_mun.values()),
                       sum(d["no_funcional"] for d in por_mun.values()),
                       sum(d["total"] for d in por_mun.values())]
            rows3.append(totals3)
            self._add_paragraph(f"Tabla. Resumen estado del parque por municipio {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers3, rows3)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        g2 = self.graficos.get("operaciones_torta")
        if g2:
            self._add_paragraph(f"Gráfico. Funcionalidad de visitas {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g2, 4.0)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

    def _cap_reposiciones(self):
        repos = self.metricas.get("reposiciones", {})
        total_comp = repos.get("total_componentes", 0)
        bat = repos.get("baterias", 0)
        ctrl = repos.get("controladores", 0)
        inv = repos.get("inversores", 0)

        self._add_paragraph(
            f"Durante el mes de {self.mes_nombre} de {self.ano}, se realizaron actividades de reposición "
            f"de equipos con el fin de restablecer la funcionalidad de los sistemas. Se instalaron en total "
            f"{fmt_num(total_comp)} componentes: {fmt_num(bat)} baterías, {fmt_num(ctrl)} controladores y "
            f"{fmt_num(inv)} inversores."
        )

        por_mun = repos.get("por_municipio", {})
        if por_mun:
            headers = ["Proyecto", "Batería", "Controlador", "Inversor", "Total"]
            rows = [[mun, d["bateria"], d["controlador"], d["inversor"], d["total"]]
                    for mun, d in por_mun.items()]
            rows.append(["Total", bat, ctrl, inv, total_comp])
            self._add_paragraph(f"Tabla. Relación de reposiciones {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        g = self.graficos.get("reposiciones")
        if g:
            self._add_image(g, 6.5)

    def _cap_gestion_comercial(self):
        self._add_paragraph(
            f"Durante el mes de {self.mes_nombre} de {self.ano}, la gestión comercial abarcó "
            "la atención al usuario en sedes, la gestión de PQR, las actividades sociales y "
            "la administración de subsidios."
        )

    def _cap_gestion_sedes(self):
        asist = self.metricas.get("asistencia", {})
        total = asist.get("total", 0)
        presencial = asist.get("presencial", 0)
        digital = asist.get("digital", 0)

        self._add_paragraph(
            f"Durante el periodo reportado se realizaron un total de {fmt_num(total)} gestiones, "
            f"de las cuales {fmt_num(presencial)} corresponden a atención presencial en sede y "
            f"{fmt_num(digital)} a contactos digitales."
        )

        por_proy = asist.get("por_proyecto", {})
        if por_proy:
            headers = ["Proyecto/Sede", "Gestiones"]
            rows = [[proy, cant] for proy, cant in por_proy.items()]
            rows.append(["TOTAL", total])
            self._add_paragraph(f"Tabla. Atención en sedes {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        g = self.graficos.get("asistencia_canal")
        if g:
            self._add_paragraph(f"Gráfico. Atención de usuarios {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g, 4.5)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

    def _cap_gestion_sac(self):
        sac = self.metricas.get("sac", {})
        total = sac.get("total", 0)
        abiertos = sac.get("abiertos", 0)
        cerrados = sac.get("cerrados", 0)
        hurtos = sac.get("hurtos", 0)

        self._add_paragraph(
            f"Durante el mes de {self.mes_nombre} de {self.ano}, se registraron {fmt_num(total)} "
            f"casos en el sistema de PQR. De estos, {fmt_num(cerrados)} fueron cerrados y "
            f"{fmt_num(abiertos)} permanecen abiertos. Se registraron además {fmt_num(hurtos)} "
            "casos relacionados con hurto de componentes."
        )

        if total > 0:
            headers = ["Estado", "Cantidad", "Porcentaje"]
            pct = lambda n: f"{n/total*100:.1f}%" if total > 0 else "0%"
            rows = [
                ["Cerrados", fmt_num(cerrados), pct(cerrados)],
                ["Abiertos", fmt_num(abiertos), pct(abiertos)],
                ["Hurtos", fmt_num(hurtos), pct(hurtos)],
                ["TOTAL", fmt_num(total), "100%"]
            ]
            self._add_paragraph(f"Tabla. Resumen PQR {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        por_tipo = sac.get("por_tipo", {})
        if por_tipo:
            headers2 = ["Tipificación", "Cantidad"]
            rows2 = [[k, v] for k, v in list(por_tipo.items())[:10]]
            self._add_paragraph(f"Tabla. PQR por tipificación {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers2, rows2)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        g1 = self.graficos.get("sac_estado")
        if g1:
            self._add_paragraph(f"Gráfico. Estado PQR – {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g1, 5.0)

        g2 = self.graficos.get("sac_tipo")
        if g2:
            self._add_paragraph(f"Gráfico. PQR por tipificación – {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g2, 6.5)

    def _cap_gestion_social(self):
        self._add_paragraph(
            f"La gestión social durante {self.mes_nombre} {self.ano} contempló actividades de "
            "relacionamiento comunitario, acompañamiento a usuarios y socialización del servicio "
            "en las zonas atendidas."
        )

    def _cap_subsidios(self):
        fact = self.metricas.get("facturacion", {})
        descuento = fact.get("total_descuento", 0)
        self._add_paragraph(
            f"Durante {self.mes_nombre} {self.ano} se aplicaron subsidios por valor de "
            f"{fmt_cop(descuento)} correspondientes a los usuarios del estrato subsidiado "
            "en las zonas no interconectadas atendidas."
        )

    def _cap_recaudo(self):
        fact = self.metricas.get("facturacion", {})
        total = fact.get("total_facturado", 0)
        neto = fact.get("neto", 0)
        usuarios = fact.get("num_usuarios", 0)

        self._add_paragraph(
            f"Durante {self.mes_nombre} {self.ano} se facturó un total de {fmt_cop(total)} "
            f"a {fmt_num(usuarios)} usuarios. El valor neto a recaudar, descontados los subsidios, "
            f"asciende a {fmt_cop(neto)}."
        )

        por_proy = fact.get("por_proyecto", {})
        if por_proy:
            headers = ["Proyecto", "Facturación (COP)", "Usuarios"]
            rows = [[p, fmt_cop(d["total"]), fmt_num(d["usuarios"])]
                    for p, d in list(por_proy.items())[:15]]
            rows.append(["TOTAL", fmt_cop(total), fmt_num(usuarios)])
            self._add_paragraph(f"Tabla. Facturación por proyecto {self.mes_nombre} {self.ano}", bold=True)
            self._build_table(headers, rows)
            self._add_paragraph("Elaboración propia de Dispower S.A.S E.S.P.", italic=True, size=8)

        g = self.graficos.get("facturacion_proyecto")
        if g:
            self._add_paragraph(f"Gráfico. Facturación por proyecto – {self.mes_nombre} {self.ano}", bold=True)
            self._add_image(g, 6.5)

    def _cap_cartera(self):
        self._add_paragraph(
            f"A continuación se presenta el análisis de cartera correspondiente al periodo "
            f"{self.mes_nombre} {self.ano}, indicando los saldos pendientes de cobro y el comportamiento "
            "del recaudo en las zonas atendidas."
        )

    def _cap_financiera(self):
        self._add_paragraph(
            f"A continuación se presenta la gestión financiera correspondiente al mes de "
            f"{self.mes_nombre} de {self.ano}."
        )

    def _cap_balance(self):
        self._add_paragraph(
            f"El balance financiero del periodo {self.mes_nombre} {self.ano} refleja la situación "
            "de ingresos, costos operativos y utilidad de la operación concesionada."
        )

    def _cap_tarifas(self):
        self._add_paragraph(
            f"Las tarifas vigentes para el periodo {self.mes_nombre} {self.ano} corresponden a las "
            "aprobadas por la CREG según el marco regulatorio aplicable a los operadores de sistemas "
            "SISFV en Zonas No Interconectadas."
        )

    def _cap_hseq(self):
        self._add_paragraph(
            f"Durante {self.mes_nombre} {self.ano} se continuó con la implementación del Sistema "
            "de Gestión en Seguridad y Salud en el Trabajo (SG-SST), dando cumplimiento a los "
            "requerimientos del Decreto 1072 de 2015 y la Resolución 0312 de 2019."
        )

    def _cap_sst(self):
        self._add_paragraph(
            "Se realizaron las actividades preventivas y de seguimiento en materia de seguridad y "
            "salud en el trabajo, incluyendo capacitaciones, inspecciones de seguridad y "
            "seguimiento a indicadores de accidentalidad."
        )

    def _cap_control_personal(self):
        self._add_paragraph(
            "A continuación se presenta el registro del personal operativo que desarrolló actividades "
            f"durante el mes de {self.mes_nombre} de {self.ano}, con la información de asistencia y "
            "cobertura de proyectos."
        )

    def _cap_ambiental(self):
        self._add_paragraph(
            f"Durante {self.mes_nombre} {self.ano} se ejecutaron las actividades de gestión ambiental "
            "contempladas en el Plan de Manejo Ambiental (PMA), incluyendo el manejo de residuos de "
            "paneles, baterías y componentes eléctricos conforme a la normativa vigente."
        )

    def _cap_actas(self):
        self._add_paragraph(
            f"A continuación se relacionan las actas de reunión celebradas durante "
            f"{self.mes_nombre} {self.ano} en el marco del contrato de concesión."
        )
        headers = ["N°", "Fecha", "Tipo de reunión", "Participantes", "Tema tratado"]
        rows = [["1", f"{self.mes_nombre} {self.ano}", "Seguimiento mensual",
                 "Dispower / Dispac / Interventoría", "Revisión indicadores del periodo"]]
        self._build_table(headers, rows)

    def _cap_juridico(self):
        self._add_paragraph(
            f"A la fecha de corte del presente informe ({self.mes_nombre} {self.ano}), "
            "el contrato de concesión se encuentra en plena vigencia y ejecución normal."
        )

    def _cap_perfeccionamiento(self):
        self._add_paragraph(
            f"El contrato de concesión {self.config.get('contrato','')} se encuentra perfeccionado "
            "y en ejecución conforme a las condiciones establecidas en el mismo."
        )

    def _cap_garantias(self):
        self._add_paragraph(
            "Las garantías contractuales se encuentran vigentes y actualizadas conforme a los "
            "requerimientos establecidos en el contrato de concesión."
        )

    def _cap_polizas_rce(self):
        self._add_paragraph(
            "Las pólizas de cumplimiento y responsabilidad civil extracontractual (RCE) "
            "se encuentran vigentes y al día con las primas correspondientes."
        )

    def _cap_poliza_todo_riesgo(self):
        self._add_paragraph(
            "La póliza todo riesgo para la infraestructura concesionada se encuentra vigente, "
            "cubriendo los equipos y sistemas SISFV operados bajo el contrato de concesión."
        )

    def _cap_fiducia(self):
        self._add_paragraph(
            f"A continuación se presenta el estado de los recursos administrados mediante "
            f"fiducia durante el periodo {self.mes_nombre} {self.ano}."
        )

    def _cap_giros_fiducia(self):
        self._add_paragraph(
            f"Durante {self.mes_nombre} {self.ano} se realizaron los giros correspondientes "
            "al contrato fiduciario de acuerdo con el cronograma establecido."
        )

    def _cap_distribucion_saldo(self):
        self._add_paragraph(
            "La distribución del saldo fiduciario al cierre del periodo corresponde a los "
            "recursos disponibles para la operación y mantenimiento del proyecto."
        )

    # ------------------------------------------------------------------ #
    # Sección de inconsistencias
    # ------------------------------------------------------------------ #
    def agregar_seccion_inconsistencias(self, inconsistencias: list):
        if not inconsistencias:
            return
        self.doc.add_page_break()
        self._add_heading("OBSERVACIONES E INCONSISTENCIAS DETECTADAS", level=1)
        self._add_paragraph(
            f"Durante el procesamiento de las bases de datos correspondientes a "
            f"{self.mes_nombre} {self.ano}, el sistema detectó las siguientes "
            f"inconsistencias que requieren atención:"
        )

        headers = ["Base de origen", "Hoja/Archivo", "NIU", "Tipo", "Descripción", "Recomendación"]
        rows = [[
            inc.get("base", ""),
            inc.get("hoja", ""),
            inc.get("nui", ""),
            inc.get("tipo", ""),
            inc.get("descripcion", ""),
            inc.get("recomendacion", "")
        ] for inc in inconsistencias]
        self._build_table(headers, rows, [1200, 1200, 800, 1200, 2400, 2400])
